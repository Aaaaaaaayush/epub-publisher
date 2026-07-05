from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches
from app.utils.logger import logger

class DocxExtractor:
    """
    Deterministic extraction engine for DOCX files.
    Reads paragraphs, headings, inline formatting (bold/italic), nested lists,
    and tables in exact document reading order, outputting clean Markdown.
    """
    
    def __init__(self, file_path: Path, media_dir: Path = None):
        self.file_path = file_path
        self.doc = Document(str(file_path))
        self.list_counters = {}
        
        # Determine media extraction path
        if media_dir is None:
            root_dir = Path(__file__).resolve().parent.parent.parent
            self.media_dir = root_dir / "data" / "extracted" / "media"
        else:
            self.media_dir = Path(media_dir)
            
        self.extract_images()
        
    def paragraph_has_images(self, p: Paragraph) -> bool:
        """Returns True if the paragraph contains any drawings or picture XML elements."""
        return bool(p._element.xpath('.//*[local-name()="drawing"] or .//*[local-name()="pict"] or .//*[local-name()="imagedata"]'))
        
    def extract_images(self):
        import zipfile
        try:
            self.media_dir.mkdir(parents=True, exist_ok=True)
            with zipfile.ZipFile(self.file_path, 'r') as z:
                for name in z.namelist():
                    if name.startswith("word/media/"):
                        img_data = z.read(name)
                        filename = Path(name).name
                        dest = self.media_dir / filename
                        with open(dest, "wb") as f:
                            f.write(img_data)
                        logger.info(f"DocxExtractor: Extracted image {filename} to {dest}")
        except Exception as e:
            logger.error(f"DocxExtractor: Failed to extract media: {e}")
            
    def iter_block_items(self):
        """
        Yield each paragraph or table in the document in reading order.
        """
        parent_elm = self.doc.element.body
        for child in parent_elm.iterchildren():
            if child.tag.endswith('p'):
                yield Paragraph(child, self.doc)
            elif child.tag.endswith('tbl'):
                yield Table(child, self.doc)

    def clean_run_text(self, text: str, is_bold: bool, is_italic: bool) -> str:
        """
        Wraps run text in markdown tags, shifting leading/trailing spaces outside.
        E.g., " word " -> " **word** " instead of "** word **".
        """
        if not text:
            return ""
        
        # Check if the text consists entirely of whitespace
        if text.isspace():
            return text
            
        # Extract leading/trailing whitespaces
        l_spaces = text[:len(text) - len(text.lstrip())]
        r_spaces = text[len(text.rstrip()):]
        stripped = text.strip()
        
        # Format the stripped content
        formatted = stripped
        if is_bold and is_italic:
            formatted = f"***{formatted}***"
        elif is_bold:
            formatted = f"**{formatted}**"
        elif is_italic:
            formatted = f"*{formatted}*"
            
        return f"{l_spaces}{formatted}{r_spaces}"

    def parse_paragraph_inline(self, p: Paragraph) -> str:
        """
        Parses bold, italic, and drawings within a paragraph, returning markdown text.
        Includes runs nested inside hyperlinks, simple fields, or content controls.
        Uses namespace-safe XPath queries to find drawings and picture elements.
        """
        from docx.text.run import Run
        text = ""
        has_drawings_in_runs = False
        
        r_elems = p._element.xpath('.//w:r')
        runs = [Run(r_elem, p) for r_elem in r_elems]
        
        for run in runs:
            # check drawings (namespace-safe: drawing, pict, imagedata)
            run_drawings = run._element.xpath('.//*[local-name()="drawing"] | .//*[local-name()="pict"] | .//*[local-name()="imagedata"]')
            run_img_text = ""
            if run_drawings:
                embed_ids = run._element.xpath('.//@*[local-name()="embed" or local-name()="id"]')
                for rId in embed_ids:
                    if rId in self.doc.part.rels:
                        rel = self.doc.part.rels[rId]
                        target = rel.target_ref
                        img_name = Path(target).name
                        run_img_text += f"![](media/{img_name})"
                        has_drawings_in_runs = True
            
            run_text = run.text
            if run_text:
                is_bold = run.bold
                is_italic = run.italic
                if run.style:
                    style_lower = run.style.name.lower()
                    if is_bold is None and ("bold" in style_lower or "strong" in style_lower):
                        is_bold = True
                    if is_italic is None and ("italic" in style_lower or "emphasis" in style_lower):
                        is_italic = True
                text += self.clean_run_text(run_text, bool(is_bold), bool(is_italic))
            
            text += run_img_text
            
        # Check if the paragraph has drawings that weren't captured in the runs
        p_drawings = p._element.xpath('.//*[local-name()="drawing"] | .//*[local-name()="pict"] | .//*[local-name()="imagedata"]')
        if p_drawings and not has_drawings_in_runs:
            embed_ids = p._element.xpath('.//@*[local-name()="embed" or local-name()="id"]')
            p_img_text = ""
            for rId in embed_ids:
                if rId in self.doc.part.rels:
                    rel = self.doc.part.rels[rId]
                    target = rel.target_ref
                    img_name = Path(target).name
                    p_img_text += f"\n\n![](media/{img_name})\n\n"
            text += p_img_text
            
        if not text and p.text:
            text = p.text
            
        return text

    def get_list_details(self, p: Paragraph) -> tuple[bool, int, str, str]:
        """
        Detects if a paragraph is a list item.
        Returns:
            is_list (bool)
            nesting_level (int) starting at 0
            list_type (str) - 'bullet' or 'number'
            prefix (str) - e.g. "- ", "1. ", "a) "
        """
        style_name = p.style.name if p.style else ""
        text = p.text.strip()
        
        # Check if the paragraph starts with a multi-level index (e.g. 1.1, 1.3.2)
        # We do NOT treat multi-level indexes as list items.
        if re.match(r'^\d+\.(?:\d+|[a-zA-Z]\b)', text):
            return False, 0, '', ''
        
        # 1. Native Word numbering detection (numPr)
        numPr = p._element.xpath('.//*[local-name()="numPr"]')
        if numPr:
            # Extract w:ilvl and w:numId val attributes safely using prefix-less xpath
            ilvl_el = p._element.xpath('.//*[local-name()="ilvl"]/@*[local-name()="val"]')
            level = int(ilvl_el[0]) if ilvl_el else 0
            
            numId_el = p._element.xpath('.//*[local-name()="numId"]/@*[local-name()="val"]')
            numId_val = numId_el[0] if numId_el else "0"
            
            list_type = 'bullet'  # default fallback
            try:
                num_part = self.doc.part.numbering_part
                if num_part:
                    nums = num_part.element.xpath(f'//*[local-name()="num" and @*[local-name()="numId"]="{numId_val}"]')
                    if nums:
                        abs_ids = nums[0].xpath('.//*[local-name()="abstractNumId"]/@*[local-name()="val"]')
                        if abs_ids:
                            abs_id = abs_ids[0]
                            abs_nums = num_part.element.xpath(f'//*[local-name()="abstractNum" and @*[local-name()="abstractNumId"]="{abs_id}"]')
                            if abs_nums:
                                numFmts = abs_nums[0].xpath(f'.//*[local-name()="lvl" and @*[local-name()="ilvl"]="{level}"]/*[local-name()="numFmt"]/@*[local-name()="val"]')
                                if numFmts:
                                    numFmt = numFmts[0].lower()
                                    if numFmt != 'bullet':
                                        list_type = 'number'
            except Exception as e:
                logger.debug(f"Failed parsing abstractNumId for numId {numId_val}: {e}")
                
            if list_type == 'number':
                key = (numId_val, level)
                if key not in self.list_counters:
                    self.list_counters[key] = 1
                else:
                    self.list_counters[key] += 1
                prefix = f"{self.list_counters[key]}. "
            else:
                prefix = "- "
                
            return True, level, list_type, prefix
            
        # 2. Fallback: Standard DOCX list styles
        is_bullet = "bullet" in style_name.lower() or "list" in style_name.lower() and "bullet" in style_name.lower()
        is_numbered = "number" in style_name.lower() or "num" in style_name.lower()
        
        # Determine level from style name (e.g. "List Bullet 2" -> Level 1)
        level = 0
        match = re.search(r'\d+', style_name)
        if match:
            level = max(0, int(match.group(0)) - 1)
            
        # If no explicit list style, check if paragraph text starts with bullet/number patterns
        # Standard bullets: •, o, *, -, ▪
        bullet_pattern = r'^([•o\*\-\▪]|-\s+)\s*(.*)'
        # Number patterns: 1., a., i., (1), 1), a)
        number_pattern = r'^(\d+\.|\([0-9a-zA-Z]+\)|[a-zA-Z]\.|\d+\)|[a-zA-Z]\))\s*(.*)'
        
        matched_prefix = ""
        bullet_match = re.match(bullet_pattern, text)
        number_match = re.match(number_pattern, text)
        
        if bullet_match:
            is_bullet = True
            matched_prefix = bullet_match.group(1)
        elif number_match:
            is_numbered = True
            matched_prefix = number_match.group(1)
            
            # Prevent false uppercase single-letter matching (like "B. Com")
            # If the prefix is a single letter (like "B." or "A)" or "A."), and there is no list style and no left indent,
            # and it is uppercase, we do NOT treat it as a list item.
            clean_prefix = matched_prefix.strip('() \t.')
            if clean_prefix.isalpha() and len(clean_prefix) == 1 and clean_prefix.isupper():
                has_list_style = any(kw in style_name.lower() for kw in ['list', 'num', 'bullet'])
                has_left_indent = False
                if p.paragraph_format.left_indent:
                    indent_val = p.paragraph_format.left_indent.inches
                    if indent_val and indent_val > 0.1:
                        has_left_indent = True
                if not has_list_style and not has_left_indent:
                    is_numbered = False
            
            # Do not treat as list item if it is an MCQ options paragraph containing multiple options separated by newlines
            if is_numbered and "\n" in text and re.search(r'\n\s*([a-zA-Z0-9]\.|\([0-9a-zA-Z]+\)|[a-zA-Z0-9]\))', text):
                is_numbered = False

            
        # Use left indent to infer list level if not set by style
        if (is_bullet or is_numbered) and level == 0 and p.paragraph_format.left_indent:
            indent_inches = p.paragraph_format.left_indent.inches
            if indent_inches:
                # Typically lists are indented in multiples of 0.25 inches
                level = int(round(indent_inches / 0.25)) - 1
                level = max(0, level)

        list_type = 'bullet' if is_bullet else ('number' if is_numbered else '')
        
        if not (is_bullet or is_numbered):
            return False, 0, '', ''
            
        if list_type == 'bullet':
            prefix = "- "
        else:
            if matched_prefix:
                prefix = f"{matched_prefix} "
            else:
                key = (style_name, level)
                if key not in self.list_counters:
                    self.list_counters[key] = 1
                else:
                    self.list_counters[key] += 1
                prefix = f"{self.list_counters[key]}. "
                
        return True, level, list_type, prefix

    def clean_list_item_text(self, text: str, list_type: str) -> str:
        """
        Strips manual bullet characters or numbering prefixes from the start of list item text,
        safely preserving any bold/italic markdown boundaries.
        """
        t = text.strip()
        if not t:
            return ""
            
        is_bold = t.startswith("**") or "**" in t[:12]
        is_italic = (t.startswith("*") and not t.startswith("**")) or ("*" in t[:6] and not "**" in t[:6])
        
        cleaned = t
        if list_type == 'bullet':
            # Remove leading bullets, spaces, asterisks, underscores
            cleaned = re.sub(r'^[\s*_\-•o▪]+\s*', '', cleaned)
        elif list_type == 'number':
            # Match number prefix mixed with spaces/asterisks/underscores/dots
            stripped_prefix = re.match(r'^[\s*_\-•o▪]*(\d+|[a-zA-Z]|[ivxIVX]+)[\s*_\-•o▪]*[.)\-]+[\s*_\-•o▪]*', cleaned)
            if stripped_prefix:
                matched_str = stripped_prefix.group(0)
                cleaned = cleaned[len(matched_str):].strip()
            else:
                cleaned = re.sub(r'^[\s*_\-•o▪]+\s*', '', cleaned)
                
        # Restore bold/italic boundaries if they were stripped
        if is_bold and not cleaned.startswith("**"):
            # Check if there is an unmatched bold tag trailing
            if cleaned.count("**") % 2 != 0:
                cleaned = f"**{cleaned}"
        elif is_italic and not cleaned.startswith("*"):
            if cleaned.count("*") % 2 != 0:
                cleaned = f"*{cleaned}"
                
        return cleaned.strip()

    def extract_metadata(self) -> tuple[str, str]:
        """
        Dynamically extracts the book title and authors from the document.
        """
        # Titles/Keywords to ignore when looking for book title
        blacklist_keywords = [
            'author', 'about', 'acknowledg', 'syllabus', 'copyright', 'content', 'index',
            'mcq', 'multiple choice', 'choice question', 'true or false', 'true/false',
            'match the', 'quiz', 'question', 'paper', 'examination', 'assessment',
            'chapter', 'module', 'unit', 'study notes'
        ]
        
        # 1. Title Extraction
        title = None
        
        def is_valid_title(text):
            low = text.lower()
            if any(kw in low for kw in blacklist_keywords):
                return False
            if len(text) < 3 or len(text) > 100:
                return False
            return True

        # Look for Title style in first 30 paragraphs
        for p in self.doc.paragraphs[:30]:
            style_name = p.style.name if p.style else ""
            text = p.text.strip()
            if not text:
                continue
            if style_name == 'Title' and is_valid_title(text):
                title = text
                break
                
        if not title:
            # Fallback: Find first Heading 1 in first 30 paragraphs
            for p in self.doc.paragraphs[:30]:
                style_name = p.style.name if p.style else ""
                text = p.text.strip()
                if not text:
                    continue
                if style_name.startswith('Heading 1') and is_valid_title(text):
                    title = text
                    break
                    
        if not title:
            # Fallback: Clean filename
            name_without_ext = self.file_path.stem
            cleaned = re.sub(r'^[a-fA-F0-9\-]{36}_?', '', name_without_ext)
            cleaned = re.sub(r'\s*\-?\s*Formatted\s*', '', cleaned, flags=re.IGNORECASE)
            title = cleaned.replace('_', ' ').replace('-', ' ').strip().title()
            
        # 2. Author Extraction
        authors = []
        author_section_found = False
        
        # Check first for paragraphs styled as Author_Name
        for p in self.doc.paragraphs[:150]:
            style_name = p.style.name if p.style else ""
            text = p.text.strip()
            if not text:
                continue
            if 'author' in style_name.lower() and 'name' in style_name.lower():
                authors.append(text)
                
        if not authors:
            # Look for lines containing Author: or By:
            author_keywords = [
                r'^(?:author|by|written\s+by)\s*:\s*(.*)',
                r'^(?:author\s+name)\s*:\s*(.*)'
            ]
            for p in self.doc.paragraphs[:100]:
                text = p.text.strip()
                if not text:
                    continue
                for kw in author_keywords:
                    m = re.match(kw, text, re.IGNORECASE)
                    if m:
                        authors.append(m.group(1).strip())
                        break
                        
        if not authors:
            # Look under Heading 1 or text matching "Author"
            for p in self.doc.paragraphs[:150]:
                text = p.text.strip()
                if not text:
                    continue
                style_name = p.style.name if p.style else ""
                
                if (style_name.startswith('Heading') or text.startswith('#')) and author_section_found:
                    clean_header = text.lstrip('#').strip().lower()
                    if clean_header not in ['author', 'about the author', 'about author', 'authors']:
                        break
                        
                clean_text = text.lstrip('#').strip().lower() if text.startswith('#') else text.lower()
                if clean_text in ['author', 'about the author', 'about author', 'authors']:
                    author_section_found = True
                    continue
                    
                if author_section_found:
                    authors.append(text)
                    break
                    
        # Remove duplicates while preserving order
        unique_authors = []
        for a in authors:
            clean_a = re.sub(r'[\s\-:,]+$', '', a).strip()
            if clean_a and clean_a not in unique_authors:
                unique_authors.append(clean_a)
                
        author_str = ", ".join(unique_authors) if unique_authors else "Educational Board"
        
        return title, author_str

    def parse_table(self, table: Table) -> str:
        """
        Parses a Word table into clean GFM (GitHub Flavored Markdown) table.
        """
        markdown_lines = []
        
        for r_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                # Extract text for cell and clean it
                cell_content = []
                for p in cell.paragraphs:
                    p_md = self.parse_paragraph_inline(p)
                    if p_md.strip():
                        cell_content.append(p_md.strip())
                cell_text = "<br />".join(cell_content)
                cell_text = cell_text.replace("\n", "<br />")
                while "<br /><br />" in cell_text:
                    cell_text = cell_text.replace("<br /><br />", "<br />")
                cell_text = cell_text.replace("|", "\\|").strip()
                row_cells.append(cell_text)
                
            markdown_lines.append("| " + " | ".join(row_cells) + " |")
            
            # Insert header separator after first row
            if r_idx == 0:
                separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                markdown_lines.append(separator)
                
        return "\n".join(markdown_lines)

    def extract_to_markdown(self) -> str:
        """
        Processes the DOCX file and returns the complete extracted Markdown text.
        """
        logger.info(f"Starting deterministic DOCX extraction for: {self.file_path.name}")
        markdown_blocks = []
        in_list = False
        active_level = 0
        active_list_type = ''
        
        for item in self.iter_block_items():
            if isinstance(item, Table):
                # Close list if we are transitioning to a table
                if in_list:
                    self.list_counters.clear()
                    markdown_blocks.append("")
                    in_list = False
                table_md = self.parse_table(item)
                markdown_blocks.append(table_md)
                markdown_blocks.append("")  # Empty line after table
                
            elif isinstance(item, Paragraph):
                style_name = item.style.name if item.style else ""
                text = item.text.strip()
                
                # Check for headings
                is_heading = False
                heading_level = 0
                if style_name.startswith("Heading"):
                    is_heading = True
                    match = re.search(r'\d+', style_name)
                    if match:
                        heading_level = int(match.group(0))
                    else:
                        heading_level = 1
                
                if is_heading:
                    if in_list:
                        self.list_counters.clear()
                        markdown_blocks.append("")
                        in_list = False
                    
                    self.list_counters.clear()
                        
                    heading_text = self.parse_paragraph_inline(item)
                    hashes = "#" * heading_level
                    markdown_blocks.append(f"{hashes} {heading_text.strip()}")
                    markdown_blocks.append("")
                    continue
                    
                # Check for lists
                is_list, level, list_type, prefix = self.get_list_details(item)
                if is_list:
                    # Normalize first list item to level 0
                    if not in_list:
                        level = 0
                    # Smart nesting: auto-lift alphabetical lists under an active numbered list to level 1
                    elif level == 0 and active_list_type == 'number' and re.match(r'^[a-zA-Z]\b', prefix.strip()):
                        level = 1
                    # Clamp level to prevent skipping levels and causing code block parsing bugs
                    elif level > active_level + 1:
                        level = active_level + 1
                        
                    in_list = True
                    active_level = level
                    active_list_type = list_type
                    item_text = self.parse_paragraph_inline(item)
                    cleaned_text = self.clean_list_item_text(item_text, list_type)
                    
                    # Ensure any internal newlines inside the list item are indented to remain within the list item
                    lines = cleaned_text.split("\n")
                    indented_lines = [lines[0]]
                    indent = "    " * level + "    "
                    for line in lines[1:]:
                        indented_lines.append(f"{indent}{line.strip()}")
                    cleaned_text = "  \n".join(indented_lines)
                    
                    # Add a blank line before list item to prevent merging with previous paragraphs/continuations
                    if markdown_blocks and markdown_blocks[-1] != "":
                        markdown_blocks.append("")
                        
                    indent = "    " * level
                    markdown_blocks.append(f"{indent}{prefix}{cleaned_text}")
                    continue
                    
                # Check for list item body continuation
                has_indent = False
                try:
                    if item.paragraph_format.left_indent is not None:
                        indent_val = item.paragraph_format.left_indent.inches
                        if indent_val and indent_val > 0.1:
                            has_indent = True
                except Exception:
                    pass
                    
                has_literal_indent = False
                if item.text and re.match(r'^[ \t]{3,}', item.text):
                    has_literal_indent = True
                    
                # Smart continuation detection:
                # A paragraph continues the current list item only if:
                #   1. It is explicitly indented (spaces/tabs before text), OR
                #   2. It starts with a known continuation keyword (Example:, Note:, etc.)
                # It does NOT continue if it is:
                #   - A bold-only line (acts as a sub-heading, e.g. **How it is done:**)
                #   - A new numbered/lettered sequence starter
                #   - An empty line (already handled above)
                is_continuation = False
                if in_list and text:
                    # Rule 1: Explicit indentation always continues
                    if has_indent or has_literal_indent:
                        is_continuation = True
                    # Rule 2: Starts with a continuation keyword
                    elif re.match(
                        r'^(example|note|notes|solution|proof|formula|rule|step|hint|tip|'
                        r'i\.e\.|e\.g\.|for example|for instance|definition|given|required|'
                        r'effects?|prevention|detection|causes?|reason)\b',
                        text, re.IGNORECASE
                    ):
                        is_continuation = True
                    # Rule 3: Bold-only text (sub-heading) → break list
                    elif re.match(r'^\*\*[^*]+\*\*\s*:?\s*$', text.strip()):
                        is_continuation = False  # sub-heading, breaks list
                    # Rule 4: Starts a new numbered/lettered item → break list
                    elif re.match(r'^\(?[0-9a-zA-Z]+[\).]\s', text.strip()):
                        is_continuation = False
                    # Default: break list if no explicit indentation or keyword
                    else:
                        is_continuation = False

                        
                if is_continuation:
                    body_text = self.parse_paragraph_inline(item)
                    # Use 4-space indent per list level to keep text inside the list item
                    # (preventing blank lines from splitting the list in CommonMark).
                    indent = "    " * (active_level + 1)
                    
                    # Prepend indent to every line of the continuation block
                    lines = body_text.split("\n")
                    indented_lines = [f"{indent}{line.strip()}" for line in lines if line.strip()]
                    formatted_continuation = "\n".join(indented_lines)
                    
                    markdown_blocks.append("")
                    markdown_blocks.append(formatted_continuation)
                    continue
                    
                # Plain paragraphs
                if in_list:
                    self.list_counters.clear()
                    markdown_blocks.append("")
                    in_list = False
                    
                # Process if it has text OR contains drawings
                if text or self.paragraph_has_images(item):
                    p_text = self.parse_paragraph_inline(item)
                    p_text = p_text.replace("\n", "  \n")
                    markdown_blocks.append(p_text.strip())
                    markdown_blocks.append("")
                    
        # Final cleanup and join
        full_markdown = "\n".join(markdown_blocks)
        # Normalize double/triple empty lines
        full_markdown = re.sub(r'\n{3,}', '\n\n', full_markdown)
        # Clean stray dollar signs wrapping acronyms/business terms, e.g. $NGOs$ -> NGOs
        full_markdown = re.sub(r'\$([a-zA-Z0-9&\-/]+)\$', r'\1', full_markdown)
        
        logger.info(f"Deterministic DOCX extraction completed for {self.file_path.name}")
        return full_markdown
