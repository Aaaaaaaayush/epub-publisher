import sys
from pathlib import Path
from docx import Document

# Add project root to path
sys.path.append(str(Path(__file__).resolve().parent.parent))

def inspect_range():
    doc_path = Path("d:/agentic_workflow/data/input/Marketing Mix-1 -Formatted.docx")
    doc = Document(str(doc_path))
    
    print("=== SEQUENTIAL PARAGRAPHS FROM INDEX 1048 TO 1070 ===")
    for idx in range(1048, min(1075, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        print(f"Index {idx}: style='{p.style.name}', text='{p.text}'")
        for r_idx, run in enumerate(p.runs):
            print(f"  Run {r_idx}: text='{run.text}'")

if __name__ == "__main__":
    inspect_range()
