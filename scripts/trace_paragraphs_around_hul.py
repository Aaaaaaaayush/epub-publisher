import docx
import sys

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    doc = docx.Document("d:/agentic_workflow/docs/Marketing Mix-1 -Formatted.docx")
    for i in range(308, 325):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            print(f"\n--- Paragraph {i} | Style: {p.style.name} ---")
            print("Raw text:", repr(p.text))
            print("Runs:")
            for j, r in enumerate(p.runs):
                print(f"  Run {j}: text={repr(r.text)}, bold={r.bold}, italic={r.italic}, style={r.style.name if r.style else None}")

if __name__ == '__main__':
    main()
