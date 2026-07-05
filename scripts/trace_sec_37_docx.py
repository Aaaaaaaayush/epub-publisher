import docx
import sys

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    doc = docx.Document("d:/agentic_workflow/docs/Marketing Mix-1 -Formatted.docx")
    
    found = False
    for i, p in enumerate(doc.paragraphs):
        if "Types of Industrial Goods" in p.text:
            found = True
            print(f"Found section at paragraph {i}")
            for j in range(max(0, i - 2), min(len(doc.paragraphs), i + 25)):
                p_curr = doc.paragraphs[j]
                print(f"\nParagraph {j} | Style: {p_curr.style.name} | Left Indent: {p_curr.paragraph_format.left_indent}")
                print("Text:", repr(p_curr.text))
                for run_idx, r in enumerate(p_curr.runs):
                    print(f"  Run {run_idx}: {repr(r.text)} | bold={r.bold} | italic={r.italic} | style={r.style.name}")
            break
            
    if not found:
        print("Types of Industrial Goods not found!")

if __name__ == '__main__':
    main()
