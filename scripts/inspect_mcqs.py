from docx import Document
from pathlib import Path

def main():
    doc_path = Path("d:/agentic_workflow/docs/Marketing Mix-1 -Formatted.docx")
    doc = Document(str(doc_path))
    
    print("=== INSPECTING MCQS AND CASE STUDIES PARAGRAPHS ===")
    
    # Let's search for "Multiple Choice Question" and "Case studies"
    mcq_start = -1
    case_start = -1
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "Multiple Choice" in text:
            mcq_start = idx
            print(f"MCQ starts at paragraph {idx}")
        if "Case studies" in text:
            case_start = idx
            print(f"Case studies starts at paragraph {idx}")
            
    if mcq_start != -1:
        print("\n--- MCQ Paragraphs Sample (10 paragraphs) ---")
        for i in range(mcq_start, min(mcq_start + 30, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            numPr = p._element.xpath('.//w:numPr')
            numId_val = "none"
            ilvl_val = "none"
            if numPr:
                numId = p._element.xpath('.//*[local-name()="numId"]/@*[local-name()="val"]')
                ilvl = p._element.xpath('.//*[local-name()="ilvl"]/@*[local-name()="val"]')
                numId_val = numId[0] if numId else "0"
                ilvl_val = ilvl[0] if ilvl else "0"
            print(f"Index {i}: text={repr(p.text[:60])}, style={repr(p.style.name)}, has_numPr={bool(numPr)}, numId={numId_val}, ilvl={ilvl_val}")

    if case_start != -1:
        print("\n--- Case Studies Paragraphs Sample (10 paragraphs) ---")
        for i in range(case_start, min(case_start + 30, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            numPr = p._element.xpath('.//w:numPr')
            numId_val = "none"
            ilvl_val = "none"
            if numPr:
                numId = p._element.xpath('.//*[local-name()="numId"]/@*[local-name()="val"]')
                ilvl = p._element.xpath('.//*[local-name()="ilvl"]/@*[local-name()="val"]')
                numId_val = numId[0] if numId else "0"
                ilvl_val = ilvl[0] if ilvl else "0"
            print(f"Index {i}: text={repr(p.text[:60])}, style={repr(p.style.name)}, has_numPr={bool(numPr)}, numId={numId_val}, ilvl={ilvl_val}")

if __name__ == "__main__":
    main()
