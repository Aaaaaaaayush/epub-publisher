import sys
from pathlib import Path
from docx import Document

# Add project root to path
sys.path.append(str(Path(__file__).resolve().parent.parent))

def safe_print(text: str):
    """Safely print strings with unencodable characters."""
    try:
        print(text)
    except UnicodeEncodeError:
        print(text.encode(sys.stdout.encoding or 'utf-8', errors='replace').decode(sys.stdout.encoding or 'utf-8'))

def inspect():
    orig_path = Path("d:/agentic_workflow/data/input/Marketing Mix-1 -Formatted.docx")
    conv_path = Path("d:/agentic_workflow/marketing_mix_principles_and_elements.docx")
    
    doc_orig = Document(str(orig_path))
    doc_conv = Document(str(conv_path))
    
    safe_print("=== INSPECTING 6.3.1.1 IN ORIGINAL ===")
    for idx, p in enumerate(doc_orig.paragraphs):
        if "6.3.1.1" in p.text or "Penetration" in p.text:
            if "Method" not in p.text and len(p.text) < 150:
                safe_print(f"Index {idx}: style='{p.style.name}', text='{p.text}'")
                for r_idx, run in enumerate(p.runs):
                    safe_print(f"  Run {r_idx}: text='{run.text}'")
                    
    safe_print("\n=== INSPECTING 6.3.1.1 IN CONVERTED ===")
    for idx, p in enumerate(doc_conv.paragraphs):
        if "6.3.1.1" in p.text or "Penetration" in p.text:
            if "Method" not in p.text and len(p.text) < 250:
                safe_print(f"Index {idx}: style='{p.style.name}', text='{p.text}'")
                for r_idx, run in enumerate(p.runs):
                    safe_print(f"  Run {r_idx}: text='{run.text}'")

if __name__ == "__main__":
    inspect()
