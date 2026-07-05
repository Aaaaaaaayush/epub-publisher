import sys
from pathlib import Path
from docx import Document

docx_path = Path("d:/agentic_workflow/data/input/Marketing Mix-1 -Formatted.docx")
doc = Document(str(docx_path))

found = False
for idx, p in enumerate(doc.paragraphs):
    if "The main objective of pricing is to:" in p.text:
        print(f"Found starting paragraph at index {idx}!")
        # print the next 20 paragraphs' text and style name
        for k in range(idx, min(len(doc.paragraphs), idx + 25)):
            p_next = doc.paragraphs[k]
            print(f"[{k}] Style='{p_next.style.name}': {repr(p_next.text)}")
        found = True
        break

if not found:
    print("Not found!")
