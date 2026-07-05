import sys
from pathlib import Path
from docx import Document

def inspect():
    docx_path = Path("d:/agentic_workflow/data/input/Marketing Mix-1 -Formatted.docx")
    doc = Document(str(docx_path))
    num_part = doc.part.numbering_part
    if not num_part:
        print("No numbering part found!")
        return
        
    print("=== PARAGRAPH NUMBERING DETAILS ===")
    for idx in range(1045, min(1080, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        numPr = p._element.xpath('.//w:numPr')
        numId_val = "none"
        ilvl_val = "none"
        numFmt = "none"
        
        if numPr:
            # Extract w:numId and w:ilvl val attributes using local-name fallback to be extremely safe
            numId = p._element.xpath('.//*[local-name()="numId"]/@*[local-name()="val"]')
            ilvl = p._element.xpath('.//*[local-name()="ilvl"]/@*[local-name()="val"]')
            
            numId_val = numId[0] if numId else "0"
            ilvl_val = ilvl[0] if ilvl else "0"
            
            # Find num element using local-name
            nums = num_part.element.xpath(f'//*[local-name()="num" and @*[local-name()="numId"]="{numId_val}"]')
            if nums:
                abs_ids = nums[0].xpath('.//*[local-name()="abstractNumId"]/@*[local-name()="val"]')
                if abs_ids:
                    abs_id = abs_ids[0]
                    # Find abstractNum element
                    abs_nums = num_part.element.xpath(f'//*[local-name()="abstractNum" and @*[local-name()="abstractNumId"]="{abs_id}"]')
                    if abs_nums:
                        # Extract numFmt val
                        numFmts = abs_nums[0].xpath(f'.//*[local-name()="lvl" and @*[local-name()="ilvl"]="{ilvl_val}"]/*[local-name()="numFmt"]/@*[local-name()="val"]')
                        if numFmts:
                            numFmt = numFmts[0]
                            
        print(f"Index {idx}: text={repr(p.text[:40])}, style={repr(p.style.name)}, has_numPr={bool(numPr)}, numId={numId_val}, ilvl={ilvl_val}, numFmt={numFmt}")

if __name__ == "__main__":
    inspect()
