from docx import Document

def inspect():
    doc = Document('data/input/Marketing Mix-1 -Formatted.docx')
    print(f"Total tables in document: {len(doc.tables)}")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx} ---")
        print(f"Rows: {len(table.rows)}")
        # Print column count of the first row
        if len(table.rows) > 0:
            print(f"Cols in row 0: {len(table.rows[0].cells)}")
        
        # Print first few cell contents
        for r_idx, row in enumerate(table.rows[:10]):
            row_texts = []
            for c_idx, cell in enumerate(row.cells):
                # Count paragraphs in this cell
                p_texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                row_texts.append(f"Col {c_idx} ({len(p_texts)} paras): {repr(p_texts[:3])}")
            print(f"  Row {r_idx}: {row_texts}")

if __name__ == "__main__":
    inspect()
