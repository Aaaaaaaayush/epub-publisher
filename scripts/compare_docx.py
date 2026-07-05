import sys
from pathlib import Path
from docx import Document

# Add root folder to path
sys.path.append(str(Path(__file__).resolve().parent.parent))

def compare_documents():
    original_path = Path("d:/agentic_workflow/data/input/Marketing Mix-1 -Formatted.docx")
    converted_path = Path("d:/agentic_workflow/marketing_mix_principles_and_elements.docx")
    
    if not original_path.exists():
        print(f"Original file not found at {original_path}")
        return
    if not converted_path.exists():
        print(f"Converted file not found at {converted_path}")
        return
        
    doc_orig = Document(str(original_path))
    doc_conv = Document(str(converted_path))
    
    print("=== DOCX STRUCTURAL COMPARISON ===")
    print(f"Original Paragraphs Count: {len(doc_orig.paragraphs)}")
    print(f"Converted Paragraphs Count: {len(doc_conv.paragraphs)}")
    print(f"Original Tables Count: {len(doc_orig.tables)}")
    print(f"Converted Tables Count: {len(doc_conv.tables)}")
    
    # 1. Compare Headings
    orig_headings = []
    for idx, p in enumerate(doc_orig.paragraphs):
        if p.style.name.startswith("Heading"):
            orig_headings.append((idx, p.style.name, p.text.strip()))
            
    conv_headings = []
    for idx, p in enumerate(doc_conv.paragraphs):
        if p.style.name.startswith("Heading"):
            conv_headings.append((idx, p.style.name, p.text.strip()))
            
    print(f"\nOriginal Headings Count: {len(orig_headings)}")
    print(f"Converted Headings Count: {len(conv_headings)}")
    
    # Check for missing headings
    print("\n--- Heading Diff ---")
    min_headings = min(len(orig_headings), len(conv_headings))
    for i in range(min_headings):
        o_style, o_text = orig_headings[i][1], orig_headings[i][2]
        c_style, c_text = conv_headings[i][1], conv_headings[i][2]
        if o_text != c_text or o_style != c_style:
            print(f"Discrepancy at index {i}:")
            print(f"  Original:  Style='{o_style}', Text='{o_text}'")
            print(f"  Converted: Style='{c_style}', Text='{c_text}'")
            
    # 2. Inspect first few paragraphs in detail
    print("\n--- First 10 Paragraphs Text Comparison ---")
    for i in range(min(15, len(doc_orig.paragraphs), len(doc_conv.paragraphs))):
        o_txt = doc_orig.paragraphs[i].text.strip()
        c_txt = doc_conv.paragraphs[i].text.strip()
        if o_txt != c_txt:
            print(f"\nParagraph {i} MISMATCH:")
            print(f"  Original : '{o_txt}'")
            print(f"  Converted: '{c_txt}'")
        else:
            print(f"Paragraph {i} MATCH: '{o_txt}'")

if __name__ == "__main__":
    compare_documents()
